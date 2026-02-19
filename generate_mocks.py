import csv
import os

from docx import Document
from openpyxl import Workbook
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


OUTPUT_DIR = "fanga_inbox"


def ensure_dir():
    os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_pdf(filename: str, lines: list[str]):
    """Create a simple PDF with text lines."""
    path = os.path.join(OUTPUT_DIR, filename)
    c = canvas.Canvas(path, pagesize=A4)
    width, height = A4
    y = height - 50
    for line in lines:
        c.drawString(50, y, line)
        y -= 20
    c.save()


def create_image(filename: str, text: str, size: tuple = (400, 300), color: str = "green"):
    """Create a simple image with text."""
    path = os.path.join(OUTPUT_DIR, filename)
    img = Image.new("RGB", size, color)
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20)
    except (OSError, IOError):
        font = ImageFont.load_default()
    draw.text((20, size[1] // 2 - 10), text, fill="white", font=font)
    img.save(path)


def generate_contrat():
    create_pdf("contrat_aissata_kone_2024.pdf", [
        "CONTRAT DE LOCATION - Moto Electrique FANGA",
        "",
        "Conductrice : Aissata Kone",
        "Date : 15/01/2024",
        "Duree : 12 mois",
        "Montant mensuel : 75 000 FCFA",
        "",
        "Conditions generales :",
        "- La conductrice s'engage a respecter le code de la route",
        "- Entretien regulier de la moto obligatoire",
        "- Restitution en bon etat a la fin du contrat",
        "",
        "Signature: ________________",
    ])


def generate_facture():
    create_pdf("facture_station_cocody_mars.pdf", [
        "FACTURE N 2024-0342",
        "",
        "Station : Cocody",
        "Mois : Mars 2024",
        "",
        "Designation          Quantite    Prix unitaire    Total",
        "Swap batterie           120        1 250 FCFA     150 000 FCFA",
        "Maintenance station       1       50 000 FCFA      50 000 FCFA",
        "",
        "TOTAL : 200 000 FCFA",
        "",
        "Mode de paiement : Virement bancaire",
    ])


def generate_photo_station():
    create_image(
        "photo_station_plateau_01.jpg",
        "Station Plateau - FANGA\nPoint de swap batterie",
        size=(400, 300),
        color="#2E7D32",
    )


def generate_rapport_xlsx():
    path = os.path.join(OUTPUT_DIR, "rapport_mensuel_conducteurs.xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "Rapport Mensuel"
    ws.append(["Nom", "Prenom", "Km parcourus", "Swaps batterie", "Mois"])
    data = [
        ("Kone", "Aissata", 1250, 45, "Mars 2024"),
        ("Traore", "Moussa", 980, 38, "Mars 2024"),
        ("Coulibaly", "Fatou", 1100, 42, "Mars 2024"),
        ("Diallo", "Ibrahim", 870, 33, "Mars 2024"),
        ("Ouattara", "Aminata", 1340, 50, "Mars 2024"),
        ("Bamba", "Seydou", 760, 28, "Mars 2024"),
        ("Toure", "Mariam", 1050, 40, "Mars 2024"),
    ]
    for row in data:
        ws.append(row)
    wb.save(path)


def generate_export_csv():
    path = os.path.join(OUTPUT_DIR, "export_transactions_fevrier.csv")
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["date", "conducteur_id", "station", "type_transaction", "montant"])
        transactions = [
            ("2024-02-01", "CD-001", "Cocody", "swap", 1250),
            ("2024-02-01", "CD-003", "Plateau", "swap", 1250),
            ("2024-02-02", "CD-005", "Yopougon", "swap", 1250),
            ("2024-02-03", "CD-002", "Cocody", "recharge", 500),
            ("2024-02-04", "CD-007", "Marcory", "swap", 1250),
            ("2024-02-05", "CD-001", "Cocody", "swap", 1250),
            ("2024-02-06", "CD-004", "Plateau", "swap", 1250),
            ("2024-02-07", "CD-006", "Treichville", "recharge", 500),
            ("2024-02-08", "CD-003", "Cocody", "swap", 1250),
            ("2024-02-09", "CD-008", "Adjame", "swap", 1250),
            ("2024-02-10", "CD-002", "Yopougon", "swap", 1250),
            ("2024-02-11", "CD-005", "Plateau", "recharge", 500),
            ("2024-02-12", "CD-001", "Marcory", "swap", 1250),
        ]
        writer.writerows(transactions)


def generate_carte_identite():
    create_image(
        "carte_identite_yacouba.png",
        "CARTE NATIONALE D'IDENTITE\nRepublique de Cote d'Ivoire\n\nYACOUBA Moussa\nNe le 12/05/1990\nAbidjan",
        size=(500, 300),
        color="#1565C0",
    )


def generate_maintenance_docx():
    path = os.path.join(OUTPUT_DIR, "maintenance_batterie_ST-002.docx")
    doc = Document()
    doc.add_heading("RAPPORT DE MAINTENANCE", level=1)
    doc.add_paragraph("Batterie : ST-002")
    doc.add_paragraph("Date : 10/03/2024")
    doc.add_paragraph("Technicien : Kouame Jean-Pierre")
    doc.add_paragraph("")
    doc.add_paragraph("Etat general : Usure normale")
    doc.add_paragraph("Capacite residuelle : 87%")
    doc.add_paragraph("Cycles de charge : 342")
    doc.add_paragraph("")
    doc.add_paragraph("Actions effectuees :")
    doc.add_paragraph("- Nettoyage des bornes de connexion")
    doc.add_paragraph("- Verification du systeme de refroidissement")
    doc.add_paragraph("- Mise a jour du firmware BMS")
    doc.add_paragraph("")
    doc.add_paragraph("Recommandation : Continuer utilisation normale. Prochain controle dans 3 mois.")
    doc.save(path)


def generate_planning():
    create_pdf("planning_equipe_avril.pdf", [
        "PLANNING EQUIPE - Avril 2024",
        "",
        "Semaine 1 (01-05 avril):",
        "  Lundi    : Equipe A - Station Cocody",
        "  Mardi    : Equipe B - Station Plateau",
        "  Mercredi : Equipe A - Station Yopougon",
        "  Jeudi    : Equipe B - Station Marcory",
        "  Vendredi : Equipe A - Station Treichville",
        "",
        "Semaine 2 (08-12 avril):",
        "  Lundi    : Equipe B - Station Cocody",
        "  Mardi    : Equipe A - Station Plateau",
        "  Mercredi : Equipe B - Station Yopougon",
        "  Jeudi    : Equipe A - Station Marcory",
        "  Vendredi : Equipe B - Station Treichville",
    ])


def generate_bon_commande():
    create_pdf("bon_de_commande_motos.pdf", [
        "BON DE COMMANDE N BC-2024-018",
        "",
        "Fournisseur : FANGA Motors CI",
        "Date : 20/02/2024",
        "",
        "Article                  Qte    Prix unitaire      Total",
        "Moto FANGA Pro            10    1 200 000 FCFA    12 000 000 FCFA",
        "Batterie supplementaire    5      150 000 FCFA       750 000 FCFA",
        "Kit maintenance            10      25 000 FCFA       250 000 FCFA",
        "",
        "TOTAL HT : 13 000 000 FCFA",
        "TVA (18%) :  2 340 000 FCFA",
        "TOTAL TTC : 15 340 000 FCFA",
    ])


def generate_screenshot_bug():
    create_image(
        "screenshot_app_bug.png",
        "FANGA App v2.1.3\n\nERROR 500\nImpossible de charger\nles transactions.\n\nVeuillez reessayer.",
        size=(300, 500),
        color="#B71C1C",
    )


def main():
    ensure_dir()
    print("Generating mock files...")

    generate_contrat()
    print("  - contrat_aissata_kone_2024.pdf")

    generate_facture()
    print("  - facture_station_cocody_mars.pdf")

    generate_photo_station()
    print("  - photo_station_plateau_01.jpg")

    generate_rapport_xlsx()
    print("  - rapport_mensuel_conducteurs.xlsx")

    generate_export_csv()
    print("  - export_transactions_fevrier.csv")

    generate_carte_identite()
    print("  - carte_identite_yacouba.png")

    generate_maintenance_docx()
    print("  - maintenance_batterie_ST-002.docx")

    generate_planning()
    print("  - planning_equipe_avril.pdf")

    generate_bon_commande()
    print("  - bon_de_commande_motos.pdf")

    generate_screenshot_bug()
    print("  - screenshot_app_bug.png")

    print(f"\nDone. {len(os.listdir(OUTPUT_DIR))} files generated in {OUTPUT_DIR}/")


if __name__ == "__main__":
    main()
